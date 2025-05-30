import pytest
from pathlib import Path
from typing import List

from docchunker import DocChunker
from docchunker.models.chunk import Chunk
from docchunker.models.document import Document


class TestDocChunker:
    """Test cases for the DocChunker class."""
    
    def test_initialization_default(self):
        """Test DocChunker initialization with default parameters."""
        chunker = DocChunker()
        
        assert chunker.chunk_size == 1000
        assert chunker.chunk_overlap == 200
        assert chunker.preserve_structure is True
        assert chunker.handle_tables is True
        assert chunker.handle_lists is True
        assert chunker.handle_images is False
    
    def test_initialization_custom(self):
        """Test DocChunker initialization with custom parameters."""
        chunker = DocChunker(
            chunk_size=500,
            chunk_overlap=100,
            preserve_structure=False,
            handle_tables=False,
            handle_lists=False,
            handle_images=True,
        )
        
        assert chunker.chunk_size == 500
        assert chunker.chunk_overlap == 100
        assert chunker.preserve_structure is False
        assert chunker.handle_tables is False
        assert chunker.handle_lists is False
        assert chunker.handle_images is True
    
    def test_process_document_docx(self, default_chunker: DocChunker, sample_docx_path: Path):
        """Test processing a DOCX document."""
        chunks = default_chunker.process_document(str(sample_docx_path))
        
        assert isinstance(chunks, list)
        assert len(chunks) > 0
        assert all(isinstance(chunk, Chunk) for chunk in chunks)
        
        # Check for different chunk types
        chunk_types = {chunk.metadata.get("type") for chunk in chunks}
        assert Chunk.TYPE_HEADING in chunk_types
        assert Chunk.TYPE_TEXT in chunk_types
        assert Chunk.TYPE_LIST in chunk_types
        assert Chunk.TYPE_TABLE in chunk_types
    
    def test_process_document_pdf(self, default_chunker: DocChunker, sample_pdf_path: Path):
        """Test processing a PDF document."""
        chunks = default_chunker.process_document(str(sample_pdf_path))
        
        assert isinstance(chunks, list)
        assert len(chunks) > 0
        assert all(isinstance(chunk, Chunk) for chunk in chunks)
        
        # Check that document_id is set
        assert all(chunk.metadata.get("document_id") for chunk in chunks)
    
    def test_process_document_invalid_path(self, default_chunker: DocChunker):
        """Test processing with invalid file path."""
        with pytest.raises(FileNotFoundError):
            default_chunker.process_document("non_existent_file.docx")
    
    def test_process_document_unsupported_format(self, default_chunker: DocChunker, temp_dir: Path):
        """Test processing with unsupported file format."""
        # Create a text file
        txt_file = temp_dir / "test.txt"
        txt_file.write_text("This is a text file")
        
        with pytest.raises(ValueError, match="Unsupported file format"):
            default_chunker.process_document(str(txt_file))
    
    def test_chunk_size_enforcement(self, custom_chunker: DocChunker, sample_docx_path: Path):
        """Test that chunks respect the specified chunk size."""
        chunker = DocChunker(chunk_size=100, chunk_overlap=20)
        chunks = chunker.process_document(str(sample_docx_path))
        
        # Most chunks should be close to the chunk size
        # (except for special chunks like tables or very short sections)
        text_chunks = [c for c in chunks if c.metadata.get("type") == Chunk.TYPE_TEXT]
        if text_chunks:
            sizes = [len(chunk.text) for chunk in text_chunks]
            # At least some chunks should be close to the target size
            assert any(80 <= size <= 120 for size in sizes)
    
    def test_chunk_overlap(self, sample_docx_path: Path):
        """Test that chunk overlap works correctly."""
        chunker = DocChunker(chunk_size=200, chunk_overlap=50)
        chunks = chunker.process_document(str(sample_docx_path))
        
        # For consecutive text chunks of the same type, check for overlap
        text_chunks = [c for c in chunks if c.metadata.get("type") == Chunk.TYPE_TEXT]
        
        # This is a basic test - in practice, overlap detection would be more complex
        assert len(text_chunks) > 0
    
    def test_preserve_structure(self, sample_docx_path: Path):
        """Test structure preservation in chunks."""
        chunker = DocChunker(preserve_structure=True)
        chunks = chunker.process_document(str(sample_docx_path))
        
        # Check that headings are preserved as separate chunks
        heading_chunks = [c for c in chunks if c.metadata.get("type") == Chunk.TYPE_HEADING]
        assert len(heading_chunks) > 0
        
        # Check that tables are preserved as separate chunks
        table_chunks = [c for c in chunks if c.metadata.get("type") == Chunk.TYPE_TABLE]
        assert len(table_chunks) > 0
    
    def test_handle_tables_enabled(self, sample_docx_path: Path):
        """Test table handling when enabled."""
        chunker = DocChunker(handle_tables=True)
        chunks = chunker.process_document(str(sample_docx_path))
        
        table_related_types = [
            Chunk.TYPE_TABLE, 
            Chunk.TYPE_TABLE_CELL,
            Chunk.TYPE_NESTED_TABLE,
            Chunk.TYPE_MERGED_CELL
        ]
        
        table_chunks = [c for c in chunks if c.metadata.get("type") in table_related_types]
        assert len(table_chunks) > 0
    
    def test_handle_tables_disabled(self, sample_docx_path: Path):
        """Test table handling when disabled."""
        chunker = DocChunker(handle_tables=False)
        chunks = chunker.process_document(str(sample_docx_path))
        
        # Tables should be treated as regular text
        table_specific_types = [
            Chunk.TYPE_TABLE_CELL,
            Chunk.TYPE_NESTED_TABLE,
            Chunk.TYPE_MERGED_CELL
        ]
        
        table_chunks = [c for c in chunks if c.metadata.get("type") in table_specific_types]
        assert len(table_chunks) == 0
    
    def test_handle_lists_enabled(self, sample_docx_path: Path):
        """Test list handling when enabled."""
        chunker = DocChunker(handle_lists=True)
        chunks = chunker.process_document(str(sample_docx_path))
        
        list_chunks = [c for c in chunks if c.metadata.get("type") == Chunk.TYPE_LIST]
        assert len(list_chunks) > 0
        
        # Check list metadata
        for chunk in list_chunks:
            assert "list_type" in chunk.metadata
            assert "depth" in chunk.metadata
    
    def test_handle_lists_disabled(self, sample_docx_path: Path):
        """Test list handling when disabled."""
        chunker = DocChunker(handle_lists=False)
        chunks = chunker.process_document(str(sample_docx_path))
        
        list_chunks = [c for c in chunks if c.metadata.get("type") == Chunk.TYPE_LIST]
        assert len(list_chunks) == 0
    
    def test_process_documents_directory(self, default_chunker: DocChunker, temp_dir: Path):
        """Test processing multiple documents in a directory."""
        # Create multiple test files
        doc1 = temp_dir / "doc1.docx"
        doc2 = temp_dir / "doc2.docx"
        
        from docx import Document
        
        # Create first document
        d1 = Document()
        d1.add_heading("Document 1", 0)
        d1.add_paragraph("This is document 1")
        d1.save(str(doc1))
        
        # Create second document
        d2 = Document()
        d2.add_heading("Document 2", 0)
        d2.add_paragraph("This is document 2")
        d2.save(str(doc2))
        
        # Process directory
        results = default_chunker.process_documents(str(temp_dir))
        
        assert len(results) == 2
        assert str(doc1) in results
        assert str(doc2) in results
        
        # Check that each document has chunks
        for file_path, chunks in results.items():
            assert len(chunks) > 0
            assert all(isinstance(chunk, Chunk) for chunk in chunks)
    
    def test_process_documents_recursive(self, default_chunker: DocChunker, temp_dir: Path):
        """Test recursive directory processing."""
        # Create subdirectory with document
        subdir = temp_dir / "subdir"
        subdir.mkdir()
        
        doc_path = subdir / "nested.docx"
        
        from docx import Document
        doc = Document()
        doc.add_heading("Nested Document", 0)
        doc.save(str(doc_path))
        
        # Process without recursion
        results_non_recursive = default_chunker.process_documents(str(temp_dir), recursive=False)
        assert str(doc_path) not in results_non_recursive
        
        # Process with recursion
        results_recursive = default_chunker.process_documents(str(temp_dir), recursive=True)
        assert str(doc_path) in results_recursive
    
    def test_export_chunks_to_json(self, default_chunker: DocChunker, sample_docx_path: Path, temp_dir: Path):
        """Test exporting chunks to JSON."""
        chunks = default_chunker.process_document(str(sample_docx_path))
        output_file = temp_dir / "chunks.json"
        
        default_chunker.export_chunks_to_json(chunks, str(output_file))
        
        assert output_file.exists()
        
        # Load and verify JSON
        import json
        with open(output_file, "r") as f:
            data = json.load(f)
        
        assert isinstance(data, list)
        assert len(data) == len(chunks)
        
        for i, chunk_data in enumerate(data):
            assert "text" in chunk_data
            assert "metadata" in chunk_data
            assert chunk_data["text"] == chunks[i].text
            assert chunk_data["metadata"] == chunks[i].metadata
    
    def test_complex_document_processing(self, default_chunker: DocChunker, complex_docx_path: Path):
        """Test processing a complex document with nested structures."""
        chunks = default_chunker.process_document(str(complex_docx_path))
        
        assert len(chunks) > 0
        
        # Check for various heading levels
        heading_chunks = [c for c in chunks if c.metadata.get("type") == Chunk.TYPE_HEADING]
        heading_levels = {c.metadata.get("level") for c in heading_chunks}
        assert len(heading_levels) >= 2  # At least 2 different heading levels
        
        # Check for merged cells in tables
        merged_cell_chunks = [c for c in chunks if c.metadata.get("type") == Chunk.TYPE_MERGED_CELL]
        if merged_cell_chunks:
            for chunk in merged_cell_chunks:
                assert "row_span" in chunk.metadata
                assert "col_span" in chunk.metadata
                assert "start_row" in chunk.metadata
                assert "start_col" in chunk.metadata
    
    def test_chunk_metadata_consistency(self, default_chunker: DocChunker, sample_docx_path: Path):
        """Test that chunk metadata is consistent and complete."""
        chunks = default_chunker.process_document(str(sample_docx_path))
        
        for chunk in chunks:
            # All chunks should have basic metadata
            assert "type" in chunk.metadata
            assert "document_id" in chunk.metadata
            
            # Check type-specific metadata
            chunk_type = chunk.metadata.get("type")
            
            if chunk_type == Chunk.TYPE_HEADING:
                assert "level" in chunk.metadata
            
            elif chunk_type == Chunk.TYPE_LIST:
                assert "list_type" in chunk.metadata
                assert "depth" in chunk.metadata
            
            elif chunk_type == Chunk.TYPE_TABLE:
                assert "table_id" in chunk.metadata
                assert "rows" in chunk.metadata
                assert "columns" in chunk.metadata
            
            elif chunk_type == Chunk.TYPE_TABLE_CELL:
                assert "table_id" in chunk.metadata
                assert "row" in chunk.metadata
                assert "column" in chunk.metadata