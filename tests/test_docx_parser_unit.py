"""Unit tests for DocxParser using small in-memory documents.

Documents are built with python-docx and parsed straight from BytesIO, so no
fixture files are needed. These cover parser paths the committed fixtures do
not reach: lists at document root (no preceding heading) and the text-based
list-marker fallback.
"""

from io import BytesIO

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docchunker.processors.docx_parser import DocxParser


def add_numbered_list_item(doc: Document, text: str, ilvl: int, num_id: int) -> None:
    """Tag a paragraph as a list item via explicit w:numPr XML."""
    paragraph = doc.add_paragraph(text)
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def parse(build_fn) -> list[dict]:
    doc = Document()
    build_fn(doc)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return DocxParser().apply(buffer)


class TestRootLevelStructures:
    def test_list_at_document_root_without_heading(self):
        def build(doc):
            add_numbered_list_item(doc, "root item one", ilvl=0, num_id=7)
            add_numbered_list_item(doc, "root item two", ilvl=0, num_id=7)

        elements = parse(build)

        assert len(elements) == 1
        container = elements[0]
        assert container["type"] == "list_container"
        assert container["num_id"] == 7
        assert [child["content"] for child in container["children"]] == [
            "root item one",
            "root item two",
        ]

    def test_paragraph_at_document_root(self):
        elements = parse(lambda doc: doc.add_paragraph("standalone paragraph"))
        assert len(elements) == 1
        assert elements[0]["type"] == "paragraph"
        assert elements[0]["content"] == "standalone paragraph"

    def test_empty_paragraphs_are_skipped(self):
        def build(doc):
            doc.add_paragraph("")
            doc.add_paragraph("   ")
            doc.add_paragraph("real content")

        elements = parse(build)
        assert len(elements) == 1
        assert elements[0]["content"] == "real content"


class TestTextBasedListFallback:
    @pytest.mark.parametrize("text", ["- dash marker item", "• dot marker item", "* star marker item"])
    def test_bullet_prefix_in_plain_paragraph_detected(self, text):
        elements = parse(lambda doc: doc.add_paragraph(text))

        assert len(elements) == 1
        container = elements[0]
        assert container["type"] == "list_container"
        assert container["num_id"] == -1
        assert container["children"][0]["content"] == text

    def test_numeric_prefix_in_plain_paragraph_detected(self):
        elements = parse(lambda doc: doc.add_paragraph("1. numbered by hand"))

        assert len(elements) == 1
        assert elements[0]["type"] == "list_container"
        assert elements[0]["children"][0]["num_id"] == -1

    def test_long_number_prefix_is_not_a_list(self):
        # "2026." has a four-digit prefix, above the parser's list threshold.
        elements = parse(lambda doc: doc.add_paragraph("2026. was quite a year"))
        assert elements[0]["type"] == "paragraph"


class TestHeadingHierarchy:
    def test_paragraph_nests_under_its_heading(self):
        def build(doc):
            doc.add_heading("Top Section", level=1)
            doc.add_paragraph("body text under section")

        elements = parse(build)
        assert len(elements) == 1
        assert elements[0]["type"] == "heading"
        assert elements[0]["level"] == 1
        children = elements[0]["children"]
        assert children[0]["type"] == "paragraph"
        assert children[0]["content"] == "body text under section"

    def test_equal_level_heading_becomes_sibling(self):
        def build(doc):
            doc.add_heading("Section A", level=1)
            doc.add_paragraph("in A")
            doc.add_heading("Section B", level=1)
            doc.add_paragraph("in B")

        elements = parse(build)
        assert [e["content"] for e in elements] == ["Section A", "Section B"]
        assert elements[0]["children"][0]["content"] == "in A"
        assert elements[1]["children"][0]["content"] == "in B"

    def test_deeper_heading_nests_under_shallower(self):
        def build(doc):
            doc.add_heading("Outer", level=1)
            doc.add_heading("Inner", level=2)
            doc.add_paragraph("innermost text")

        elements = parse(build)
        assert len(elements) == 1
        inner = elements[0]["children"][0]
        assert inner["type"] == "heading"
        assert inner["level"] == 2
        assert inner["children"][0]["content"] == "innermost text"


class TestTableParsing:
    def test_first_row_becomes_header_rest_become_data_rows(self):
        def build(doc):
            table = doc.add_table(rows=3, cols=2)
            for col, text in enumerate(["ColA", "ColB"]):
                table.cell(0, col).text = text
            table.cell(1, 0).text = "a1"
            table.cell(1, 1).text = "b1"
            table.cell(2, 0).text = "a2"
            table.cell(2, 1).text = "b2"

        elements = parse(build)
        tables = [e for e in elements if e["type"] == "table"]
        assert len(tables) == 1
        table = tables[0]
        assert table["header"] == ["ColA", "ColB"]
        assert table["data_rows"] == [["a1", "b1"], ["a2", "b2"]]
        assert table["num_rows"] == 3
        assert table["num_cols"] == 2

    def test_fully_empty_data_rows_are_dropped(self):
        def build(doc):
            table = doc.add_table(rows=3, cols=2)
            table.cell(0, 0).text = "H1"
            table.cell(0, 1).text = "H2"
            # Row 1 left entirely empty; row 2 partially filled.
            table.cell(2, 0).text = "kept"

        elements = parse(build)
        table = next(e for e in elements if e["type"] == "table")
        assert table["data_rows"] == [["kept", ""]]
