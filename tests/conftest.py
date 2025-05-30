"""
Configuration and fixtures for pytest tests.
"""

import os
import pytest
from pathlib import Path
from typing import Generator
import tempfile
import shutil

from docx import Document
from docx.shared import Inches
import pypdf
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

from docchunker import DocChunker


@pytest.fixture
def test_data_dir() -> Path:
    """Get the path to test data directory."""
    return Path(__file__).parent / "test_data"


@pytest.fixture
def temp_dir() -> Generator[Path, None, None]:
    """Create a temporary directory for test files."""
    temp_path = Path(tempfile.mkdtemp())
    yield temp_path
    shutil.rmtree(temp_path)


@pytest.fixture
def sample_docx_path(temp_dir: Path) -> Path:
    """Create a sample DOCX file for testing."""
    doc = Document()
    
    # Add title
    doc.add_heading('Test Document', 0)
    
    # Add heading
    doc.add_heading('Introduction', level=1)
    
    # Add paragraph
    doc.add_paragraph('This is a test document for the DocChunker library.')
    
    # Add another heading
    doc.add_heading('Features', level=1)
    
    # Add a list
    doc.add_paragraph('Key features:', style='List Bullet')
    doc.add_paragraph('Structure-aware chunking', style='List Bullet')
    doc.add_paragraph('Complex element handling', style='List Bullet')
    doc.add_paragraph('Multi-format support', style='List Bullet')
    
    # Add a numbered list
    doc.add_heading('Steps', level=2)
    doc.add_paragraph('First step', style='List Number')
    doc.add_paragraph('Second step', style='List Number')
    doc.add_paragraph('Third step', style='List Number')
    
    # Add a table
    doc.add_heading('Data Table', level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Name'
    header_cells[1].text = 'Type'
    header_cells[2].text = 'Value'
    
    # Add data rows
    row1 = table.rows[1].cells
    row1[0].text = 'Chunk Size'
    row1[1].text = 'Integer'
    row1[2].text = '1000'
    
    row2 = table.rows[2].cells
    row2[0].text = 'Overlap'
    row2[1].text = 'Integer'
    row2[2].text = '200'
    
    # Add more content
    doc.add_paragraph('\nThis is some text after the table.')
    
    # Save document
    file_path = temp_dir / "test_document.docx"
    doc.save(str(file_path))
    
    return file_path


@pytest.fixture
def complex_docx_path(temp_dir: Path) -> Path:
    """Create a complex DOCX file with nested structures."""
    doc = Document()
    
    # Title
    doc.add_heading('Complex Test Document', 0)
    
    # Multi-level headings
    doc.add_heading('Chapter 1: Introduction', level=1)
    doc.add_heading('1.1 Background', level=2)
    doc.add_paragraph('This section provides background information.')
    
    doc.add_heading('1.2 Objectives', level=2)
    doc.add_paragraph('The main objectives are:')
    
    # Nested lists
    doc.add_paragraph('Primary goals:', style='List Bullet')
    doc.add_paragraph('    • Improve performance', style='List Bullet')
    doc.add_paragraph('    • Enhance usability', style='List Bullet')
    doc.add_paragraph('Secondary goals:', style='List Bullet')
    doc.add_paragraph('    • Reduce costs', style='List Bullet')
    doc.add_paragraph('    • Increase adoption', style='List Bullet')
    
    # Complex table with merged cells
    doc.add_heading('Chapter 2: Data Analysis', level=1)
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Merge cells in first row
    a = table.cell(0, 0)
    b = table.cell(0, 1)
    a.merge(b)
    a.text = 'Merged Header'
    
    table.cell(0, 2).text = 'Col 3'
    table.cell(0, 3).text = 'Col 4'
    
    # Add data
    for i in range(1, 4):
        for j in range(4):
            table.cell(i, j).text = f'Cell {i},{j}'
    
    # Long paragraph
    doc.add_heading('Chapter 3: Discussion', level=1)
    long_text = ' '.join(['This is a long paragraph.'] * 50)
    doc.add_paragraph(long_text)
    
    # Nested table (table within table)
    doc.add_heading('Chapter 4: Nested Structures', level=1)
    outer_table = doc.add_table(rows=2, cols=2)
    outer_table.style = 'Light Grid Accent 1'
    
    # Add text to first cell
    outer_table.cell(0, 0).text = 'Outer cell 1'
    
    # Add a nested table to second cell
    cell = outer_table.cell(0, 1)
    # Note: python-docx doesn't directly support nested tables, 
    # so we'll simulate with paragraph formatting
    cell.text = 'Nested content here'
    
    # Save document
    file_path = temp_dir / "complex_document.docx"
    doc.save(str(file_path))
    
    return file_path


@pytest.fixture
def sample_pdf_path(temp_dir: Path) -> Path:
    """Create a sample PDF file for testing."""
    if not REPORTLAB_AVAILABLE:
        pytest.skip("reportlab not available for PDF generation")
    
    file_path = temp_dir / "test_document.pdf"
    
    # Create PDF with reportlab
    c = canvas.Canvas(str(file_path), pagesize=letter)
    
    # Add content
    c.setFont("Helvetica-Bold", 16)
    c.drawString(100, 750, "Test PDF Document")
    
    c.setFont("Helvetica-Bold", 14)
    c.drawString(100, 700, "Introduction")
    
    c.setFont("Helvetica", 12)
    c.drawString(100, 670, "This is a test PDF document for the DocChunker library.")
    
    c.setFont("Helvetica-Bold", 14)
    c.drawString(100, 620, "Features")
    
    c.setFont("Helvetica", 12)
    c.drawString(100, 590, "• Structure-aware chunking")
    c.drawString(100, 570, "• Complex element handling")
    c.drawString(100, 550, "• Multi-format support")
    
    c.setFont("Helvetica-Bold", 14)
    c.drawString(100, 500, "Data Table")
    
    # Simple table representation
    c.setFont("Helvetica", 12)
    c.drawString(100, 470, "Name | Type | Value")
    c.drawString(100, 450, "Chunk Size | Integer | 1000")
    c.drawString(100, 430, "Overlap | Integer | 200")
    
    c.showPage()
    c.save()
    
    return file_path


@pytest.fixture
def default_chunker() -> DocChunker:
    """Create a DocChunker instance with default settings."""
    return DocChunker()


@pytest.fixture
def custom_chunker() -> DocChunker:
    """Create a DocChunker instance with custom settings."""
    return DocChunker(
        chunk_size=500,
        chunk_overlap=100,
        preserve_structure=True,
        handle_tables=True,
        handle_lists=True,
        handle_images=False,
    )