"""Unit tests for rule-based heading inference in DocxParser.

Heading inference only applies to documents WITHOUT styled headings, where
authors fake headings using bold runs, larger fonts, manual numbering
("2.3.1 Title") or ALL-CAPS lines. These tests build small documents in
memory and assert:

- positive signals (bold, larger font, all-caps, numbering depth) produce
  heading nodes with deterministic levels;
- negative signals (plain short lines, inline bold emphasis, trailing
  punctuation, list items) never produce headings;
- styled documents are completely unaffected (byte-identical parses), both
  for in-memory documents and for every committed DOCX fixture.
"""

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from docchunker.processors.docx_parser import DocxParser

FIXTURES_DIR = Path(__file__).parent.parent / "data" / "unittests"


def add_numbered_list_item(doc: Document, text: str, ilvl: int, num_id: int) -> None:
    """Tag a paragraph as a list item via explicit w:numPr XML."""
    paragraph = doc.add_paragraph(text)
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def add_formatted_paragraph(doc: Document, text: str, bold: bool = False, size_pt: float | None = None):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    return paragraph


def parse(build_fn, **parser_kwargs) -> list[dict]:
    doc = Document()
    build_fn(doc)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return DocxParser(**parser_kwargs).apply(buffer)


BODY_TEXT = (
    "This body paragraph provides enough surrounding prose to look like "
    "ordinary document content rather than a heading."
)


class TestPositiveInference:
    def test_bold_short_paragraph_becomes_heading(self):
        def build(doc):
            add_formatted_paragraph(doc, "Overview", bold=True)
            doc.add_paragraph(BODY_TEXT)

        elements = parse(build)
        assert elements[0]["type"] == "heading"
        assert elements[0]["level"] == 1
        assert elements[0]["content"] == "Overview"
        assert elements[0]["children"][0]["type"] == "paragraph"

    def test_all_caps_short_line_becomes_heading(self):
        def build(doc):
            add_formatted_paragraph(doc, "GENERAL PROVISIONS")
            doc.add_paragraph(BODY_TEXT)

        elements = parse(build)
        assert elements[0]["type"] == "heading"
        assert elements[0]["content"] == "GENERAL PROVISIONS"

    def test_font_size_tiers_map_to_levels(self):
        def build(doc):
            add_formatted_paragraph(doc, "Alpha Section", bold=True, size_pt=18)
            doc.add_paragraph(BODY_TEXT)
            add_formatted_paragraph(doc, "Beta Subsection", bold=True, size_pt=14)
            doc.add_paragraph(BODY_TEXT)
            add_formatted_paragraph(doc, "Gamma Bold Only", bold=True)
            doc.add_paragraph(BODY_TEXT)

        elements = parse(build)
        alpha = elements[0]
        assert (alpha["type"], alpha["level"]) == ("heading", 1)
        beta = alpha["children"][1]
        assert (beta["type"], beta["level"]) == ("heading", 2)
        # Same-size bold-only: one level deeper than the deepest size tier.
        gamma = beta["children"][1]
        assert (gamma["type"], gamma["level"]) == ("heading", 3)

    def test_numbering_prefix_depth_refines_level(self):
        def build(doc):
            add_formatted_paragraph(doc, "2. Requirements", bold=True)
            doc.add_paragraph(BODY_TEXT)
            add_formatted_paragraph(doc, "2.3 Interfaces", bold=True)
            doc.add_paragraph(BODY_TEXT)
            add_formatted_paragraph(doc, "2.3.1 Wire Protocol", bold=True)
            doc.add_paragraph(BODY_TEXT)

        elements = parse(build)
        top = elements[0]
        assert (top["type"], top["level"]) == ("heading", 1)
        mid = top["children"][1]
        assert (mid["type"], mid["level"]) == ("heading", 2)
        deep = mid["children"][1]
        assert (deep["type"], deep["level"]) == ("heading", 3)
        assert deep["children"][0]["type"] == "paragraph"


class TestNegativeInference:
    def test_short_plain_paragraph_is_not_a_heading(self):
        """A lone short paragraph without any emphasis must stay body text."""
        def build(doc):
            doc.add_paragraph("Just a short line")
            doc.add_paragraph(BODY_TEXT)

        elements = parse(build)
        assert [e["type"] for e in elements] == ["paragraph", "paragraph"]

    def test_inline_bold_emphasis_in_long_paragraph_is_not_a_heading(self):
        def build(doc):
            paragraph = doc.add_paragraph("The committee decided that ")
            run = paragraph.add_run("immediate escalation")
            run.bold = True
            paragraph.add_run(
                " applies only when the incident affects more than one region"
            )

        elements = parse(build)
        assert len(elements) == 1
        assert elements[0]["type"] == "paragraph"

    def test_bold_short_paragraph_ending_with_period_is_not_a_heading(self):
        def build(doc):
            add_formatted_paragraph(doc, "This is emphasized.", bold=True)
            doc.add_paragraph(BODY_TEXT)

        elements = parse(build)
        assert all(e["type"] == "paragraph" for e in elements)

    def test_fully_bold_long_paragraph_is_not_a_heading(self):
        def build(doc):
            add_formatted_paragraph(doc, BODY_TEXT, bold=True)

        elements = parse(build)
        assert elements[0]["type"] == "paragraph"

    def test_bold_list_item_stays_a_list_item(self):
        def build(doc):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("Bold checklist entry")
            run.bold = True
            p_pr = paragraph._p.get_or_add_pPr()
            num_pr = OxmlElement("w:numPr")
            ilvl_el = OxmlElement("w:ilvl")
            ilvl_el.set(qn("w:val"), "0")
            num_id_el = OxmlElement("w:numId")
            num_id_el.set(qn("w:val"), "9")
            num_pr.append(ilvl_el)
            num_pr.append(num_id_el)
            p_pr.append(num_pr)

        elements = parse(build)
        assert elements[0]["type"] == "list_container"
        assert elements[0]["children"][0]["content"] == "Bold checklist entry"

    def test_all_bold_document_triggers_ratio_guard(self):
        """When most paragraphs look like headings, infer nothing."""
        def build(doc):
            for i in range(10):
                add_formatted_paragraph(doc, f"Bold short line {i}", bold=True)

        elements = parse(build)
        assert all(e["type"] == "paragraph" for e in elements)


class TestStyledDocumentsUnchanged:
    def test_styled_document_ignores_fake_headings(self):
        """One styled heading anywhere disables inference for the whole doc."""
        def build(doc):
            doc.add_heading("Real Styled Heading", level=1)
            doc.add_paragraph(BODY_TEXT)
            add_formatted_paragraph(doc, "Fake Bold Heading", bold=True, size_pt=16)
            doc.add_paragraph(BODY_TEXT)

        with_inference = parse(build)
        without_inference = parse(build, infer_headings=False)
        assert with_inference == without_inference

        heading = with_inference[0]
        assert heading["content"] == "Real Styled Heading"
        # The bold paragraph stayed a plain paragraph under the styled heading.
        assert [c["type"] for c in heading["children"]] == [
            "paragraph", "paragraph", "paragraph",
        ]

    @pytest.mark.parametrize(
        "fixture_name",
        sorted(p.name for p in FIXTURES_DIR.glob("*.docx")),
    )
    def test_committed_fixtures_parse_identically(self, fixture_name):
        """Regression gate: inference must not alter any committed fixture
        that uses heading styles; for the unstyled fixture it only adds
        heading nodes when enabled (covered by the YAML suite)."""
        fixture_path = FIXTURES_DIR / fixture_name
        default_parse = DocxParser().apply(str(fixture_path))
        if fixture_name == "unstyled_headings.docx":
            headings = [e for e in default_parse if e["type"] == "heading"]
            assert headings, "unstyled fixture should gain inferred headings"
            return
        assert default_parse == DocxParser(infer_headings=False).apply(str(fixture_path))

    def test_opt_out_flag_disables_inference(self):
        def build(doc):
            add_formatted_paragraph(doc, "Overview", bold=True)
            doc.add_paragraph(BODY_TEXT)

        elements = parse(build, infer_headings=False)
        assert all(e["type"] == "paragraph" for e in elements)
