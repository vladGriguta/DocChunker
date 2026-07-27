"""Generate DOCX test fixtures for the YAML-driven unit tests.

Run from the repository root:

    python data/unittests/generate_fixtures.py

Regenerates:
    data/unittests/complex_tables.docx
    data/unittests/complex_lists.docx
    data/unittests/unstyled_headings.docx

The generated files are committed to the repository so tests do not depend on
this script at runtime, but the script is kept so fixtures stay reproducible.

No PDF counterparts are generated: there is no offline DOCX->PDF conversion
dependency in this project, and the PDF test runner automatically skips YAML
configs without a matching .pdf file.
"""

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

OUTPUT_DIR = Path(__file__).parent


def add_list_item(doc: Document, text: str, ilvl: int, num_id: int):
    """Add a paragraph tagged as a list item via explicit numbering XML.

    DocxParser reads w:numPr (w:ilvl + w:numId) directly from each paragraph's
    properties, so setting them explicitly is the most reliable way to produce
    nested lists programmatically.
    """
    paragraph = doc.add_paragraph(text)
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)
    return paragraph


def fill_table(table, header: list[str], rows: list[list[str]]) -> None:
    for col_idx, text in enumerate(header):
        table.cell(0, col_idx).text = text
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, text in enumerate(row):
            table.cell(row_idx, col_idx).text = text


def build_complex_tables(path: Path) -> None:
    doc = Document()
    doc.add_heading("Complex Tables Test Document", level=1)
    doc.add_paragraph(
        "This document exercises table edge cases: merged cells, wide tables, "
        "header-only tables and large tables that must be split across chunks."
    )

    # --- Section 1: merged cells -------------------------------------------
    doc.add_heading("Merged Cells Table", level=2)
    doc.add_paragraph("The table below contains both vertical and horizontal merges.")
    table = doc.add_table(rows=5, cols=3)
    fill_table(
        table,
        header=["Region", "Quarter", "Revenue"],
        rows=[
            ["EMEA shared region", "Q1", "Revenue was 1200 units"],
            ["", "Q2", "Revenue was 1350 units"],
            ["Full year summary spanning all quarter columns", "", ""],
            ["APAC region", "Q1", "Revenue was 800 units"],
        ],
    )
    # Vertical merge: EMEA spans two data rows.
    table.cell(1, 0).merge(table.cell(2, 0))
    # Horizontal merge: summary row spans all three columns.
    table.cell(3, 0).merge(table.cell(3, 2))

    # --- Section 2: wide table ---------------------------------------------
    doc.add_heading("Wide Table", level=2)
    doc.add_paragraph("A table with twelve columns to test wide-row formatting.")
    n_cols = 12
    wide = doc.add_table(rows=3, cols=n_cols)
    fill_table(
        wide,
        header=[f"Metric{i:02d}" for i in range(1, n_cols + 1)],
        rows=[
            [f"wide-r1-c{i:02d}" for i in range(1, n_cols + 1)],
            [f"wide-r2-c{i:02d}" for i in range(1, n_cols + 1)],
        ],
    )

    # --- Section 3: header-only table --------------------------------------
    doc.add_heading("Header Only Table", level=2)
    doc.add_paragraph("The following table consists of a single header row.")
    header_only = doc.add_table(rows=1, cols=3)
    fill_table(
        header_only,
        header=["Status Code", "Status Meaning", "Escalation Contact"],
        rows=[],
    )

    # --- Section 4: large table forcing multi-chunk splitting ---------------
    doc.add_heading("Large Product Catalog", level=2)
    doc.add_paragraph("A long table that must be split into several chunks.")
    n_rows = 30
    catalog = doc.add_table(rows=n_rows + 1, cols=3)
    fill_table(
        catalog,
        header=["Product Name", "Description", "Price"],
        rows=[
            [
                f"Gadget model {i:02d}",
                f"Catalog entry number {i:02d} describing a durable industrial "
                f"gadget with extended warranty and modular accessories.",
                f"{i * 10 + 99} USD",
            ]
            for i in range(1, n_rows + 1)
        ],
    )

    doc.save(str(path))


def build_complex_lists(path: Path) -> None:
    doc = Document()
    doc.add_heading("Complex Lists Test Document", level=1)
    doc.add_paragraph(
        "This document exercises list edge cases: deep nesting, mixed list "
        "types, lists interrupted by paragraphs and lists following tables."
    )

    # --- Section 1: deep nesting (5 levels) ---------------------------------
    doc.add_heading("Deep Nesting", level=2)
    doc.add_paragraph("The plan below is nested five levels deep.")
    add_list_item(doc, "Deploy platform level one", ilvl=0, num_id=100)
    add_list_item(doc, "Provision infrastructure level two", ilvl=1, num_id=100)
    add_list_item(doc, "Configure network level three", ilvl=2, num_id=100)
    add_list_item(doc, "Open firewall ports level four", ilvl=3, num_id=100)
    add_list_item(doc, "Whitelist monitoring agents level five", ilvl=4, num_id=100)
    add_list_item(doc, "Rotate credentials level four sibling", ilvl=3, num_id=100)
    add_list_item(doc, "Configure storage level three sibling", ilvl=2, num_id=100)
    add_list_item(doc, "Verify deployment level one sibling", ilvl=0, num_id=100)

    # --- Section 2: mixed numbered and bulleted lists -----------------------
    doc.add_heading("Mixed List Types", level=2)
    doc.add_paragraph("A numbered checklist followed by a bulleted glossary.")
    add_list_item(doc, "Numbered step alpha: collect requirements", ilvl=0, num_id=200)
    add_list_item(doc, "Numbered substep alpha-one: interview users", ilvl=1, num_id=200)
    add_list_item(doc, "Numbered substep alpha-two: draft specification", ilvl=1, num_id=200)
    add_list_item(doc, "Numbered step beta: review budget", ilvl=0, num_id=200)
    # Style-based bullets (no explicit numbering XML) exercise the parser's
    # style-name fallback detection path.
    doc.add_paragraph("Bullet term latency: time before first byte", style="List Bullet")
    doc.add_paragraph("Bullet term throughput: bytes moved per second", style="List Bullet")

    # --- Section 3: list interrupted by a paragraph -------------------------
    doc.add_heading("Interrupted List", level=2)
    add_list_item(doc, "Interrupted item one: unplug the appliance", ilvl=0, num_id=300)
    add_list_item(doc, "Interrupted item two: remove the side panel", ilvl=0, num_id=300)
    doc.add_paragraph(
        "Warning interlude: always discharge the capacitor before continuing "
        "with the remaining steps."
    )
    add_list_item(doc, "Interrupted item three: replace the filter", ilvl=0, num_id=300)
    add_list_item(doc, "Interrupted item four: reattach the side panel", ilvl=0, num_id=300)

    # --- Section 4: list immediately following a table ----------------------
    doc.add_heading("List After Table", level=2)
    table = doc.add_table(rows=3, cols=2)
    fill_table(
        table,
        header=["Ingredient", "Amount"],
        rows=[
            ["Flour for the base", "500 grams"],
            ["Water at room temperature", "300 millilitres"],
        ],
    )
    add_list_item(doc, "After-table step one: mix the ingredients", ilvl=0, num_id=400)
    add_list_item(doc, "After-table step two: rest the dough", ilvl=0, num_id=400)
    add_list_item(doc, "After-table step three: bake at high heat", ilvl=0, num_id=400)

    doc.save(str(path))


def add_formatted_paragraph(
    doc: Document, text: str, bold: bool = False, size_pt: float | None = None
):
    """Add a plain (non-styled) paragraph with optional run-level formatting.

    Used to fake headings the way real-world authors do: bold runs, larger
    fonts and manual numbering instead of Word heading styles.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    return paragraph


def build_unstyled_headings(path: Path) -> None:
    """Document with NO heading styles: headings are faked via formatting.

    Exercises all three heading-inference signal types across three levels:
    - bold + larger font + numbering  -> level 1 ("1. Introduction")
    - bold + numbering depth 2/3      -> levels 2/3 ("1.1 Purpose", "2.2.1 ...")
    - bold only (body size)           -> one level below the size tiers
    - ALL-CAPS only                   -> one level below the size tiers
    Plus negative cases: short plain paragraphs and inline bold emphasis
    inside a long paragraph must NOT become headings.
    """
    doc = Document()

    # Lead body paragraph before any heading (ends with '.').
    doc.add_paragraph(
        "Atlas is an internal telemetry platform. This overview document "
        "describes its architecture and operating procedures for new engineers."
    )

    # --- Section 1: bold + 16pt + numbering -> level 1 ----------------------
    add_formatted_paragraph(doc, "1. Introduction", bold=True, size_pt=16)
    doc.add_paragraph(
        "The introduction section explains why Atlas exists and who should "
        "read this guide before touching production systems."
    )

    add_formatted_paragraph(doc, "1.1 Purpose", bold=True)
    doc.add_paragraph(
        "Atlas collects telemetry from edge devices and stores it for later "
        "analysis by the data science team."
    )
    # Negative case: short plain paragraph, no emphasis -> must stay body text.
    doc.add_paragraph("Refer to the glossary for terminology")
    doc.add_paragraph(
        "Feedback on this specification should go to the platform working group."
    )

    add_formatted_paragraph(doc, "1.2 Scope", bold=True)
    # Negative case: bold inline emphasis inside a long paragraph.
    scope = doc.add_paragraph(
        "The scope of this specification covers ingestion, storage and reporting. "
    )
    emphasis = scope.add_run("Real-time alerting")
    emphasis.bold = True
    scope.add_run(
        " is explicitly out of scope for the first release and will follow in "
        "a later revision of this document."
    )

    # --- Section 2: nested numbering + a list and a table -------------------
    add_formatted_paragraph(doc, "2. System Architecture", bold=True, size_pt=16)
    doc.add_paragraph(
        "The platform is composed of loosely coupled services communicating "
        "over a shared message bus."
    )

    add_formatted_paragraph(doc, "2.1 Components", bold=True)
    doc.add_paragraph("The services below make up the core of the platform.")
    add_list_item(doc, "Collector service accepts device payloads", ilvl=0, num_id=500)
    add_list_item(doc, "Registry service tracks device identity", ilvl=0, num_id=500)
    add_list_item(doc, "Reporting service renders dashboards", ilvl=0, num_id=500)

    add_formatted_paragraph(doc, "2.2 Data Flow", bold=True)
    doc.add_paragraph(
        "Telemetry moves through the system in three stages described below."
    )

    add_formatted_paragraph(doc, "2.2.1 Ingestion Pipeline", bold=True)
    doc.add_paragraph(
        "During ingestion the collector validates each payload and appends it "
        "to the durable event log for downstream consumers."
    )

    # --- Bold-only heading (no numbering, body font size) -------------------
    add_formatted_paragraph(doc, "Deployment Checklist", bold=True)
    doc.add_paragraph("Every release must pass the checklist before rollout.")
    table = doc.add_table(rows=3, cols=2)
    fill_table(
        table,
        header=["Environment", "Approver"],
        rows=[
            ["Staging cluster", "Platform lead"],
            ["Production cluster", "Release manager"],
        ],
    )

    # --- ALL-CAPS heading (no bold, body font size) -------------------------
    add_formatted_paragraph(doc, "APPENDIX A GLOSSARY")
    doc.add_paragraph(
        "Telemetry means measurements emitted by devices without operator "
        "interaction."
    )

    doc.save(str(path))


def main() -> None:
    build_complex_tables(OUTPUT_DIR / "complex_tables.docx")
    build_complex_lists(OUTPUT_DIR / "complex_lists.docx")
    build_unstyled_headings(OUTPUT_DIR / "unstyled_headings.docx")
    print(f"Fixtures written to {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
