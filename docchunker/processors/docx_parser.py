import logging
import re
from collections import Counter
from typing import Any, Union, BinaryIO
import docx
from docx.table import Table
from docx.text.paragraph import Paragraph
import docx.document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

logger = logging.getLogger(__name__)


class DocxParser:
    """Parses DOCX to a hierarchical structure of elements.

    Heading detection works in two layers:

    1. **Styled headings** (``Heading 1``, ``Heading 2``, ...) are always used
       when present. They take absolute precedence: if a document contains any
       paragraph with a ``Heading N`` style, no heading inference is performed
       and the output is identical to inference-free parsing.
    2. **Inferred headings** (see :meth:`_infer_heading_levels`) kick in only
       for documents that use no heading styles at all — the common case of
       authors faking headings with bold text, larger fonts, ALL-CAPS lines or
       "2.3.1 Title"-style numbering. Inference is deliberately conservative:
       false negatives merely lose some context, while a false positive would
       corrupt the heading context of everything below it.
    """

    # ---- Heading inference thresholds (conservative on purpose) -----------
    #: Inferred headings must be short.
    MAX_INFERRED_HEADING_LENGTH = 80
    #: Fraction of run characters that must be bold for "dominantly bold".
    MIN_BOLD_RATIO = 0.8
    #: Minimum number of letters for the ALL-CAPS signal to apply.
    MIN_ALL_CAPS_LETTERS = 4
    #: Heading levels are capped at this depth.
    MAX_HEADING_LEVEL = 6
    #: Fallback body font size (Word default) when none can be measured.
    DEFAULT_BODY_FONT_SIZE_PT = 11.0
    #: If more than this fraction of paragraphs qualify as headings, the
    #: emphasis signals are considered unreliable and nothing is inferred.
    MAX_CANDIDATE_RATIO = 0.5
    #: The ratio guard above only applies to documents at least this long.
    MIN_PARAGRAPHS_FOR_RATIO_GUARD = 8

    #: A heading must not end like a sentence or list-ish fragment.
    SENTENCE_END_CHARS = ('.', ',', ';', ':', '!', '?')
    #: Textual bullet markers disqualify a paragraph from heading inference.
    BULLET_PREFIXES = ('- ', '• ', '* ')
    #: Matches "1 Title", "2. Title", "2.3 Title", "2.3.1 Title", ...
    NUMBERING_PREFIX_RE = re.compile(r"^(\d+(?:\.\d+)*)\.?\s+\S")

    def __init__(self, infer_headings: bool = True):
        """
        Args:
            infer_headings: When True (default), infer headings from formatting
                in documents that contain no styled headings. Styled documents
                are never affected by this flag.
        """
        self.current_heading_level = 0
        self.infer_headings = infer_headings
        self._inferred_heading_levels: dict[CT_P, int] = {}

    def apply(self, file_input: Union[str, BinaryIO]) -> list[dict[str, Any]]:
        """Parse DOCX and return a hierarchical list of element dictionaries.

        Args:
            file_input: Either a file path (str) or a file-like object (BinaryIO)
        """
        doc = docx.Document(file_input)
        self.current_heading_level = 0
        self._inferred_heading_levels = (
            self._infer_heading_levels(doc) if self.infer_headings else {}
        )

        hierarchical_elements = self._parse_content_elements(doc)
        return hierarchical_elements

    # ------------------------------------------------------------------
    # Heading inference for unstyled documents
    # ------------------------------------------------------------------

    def _infer_heading_levels(self, doc: docx.document.Document) -> dict[CT_P, int]:
        """Infer heading levels for documents that use no heading styles.

        Rules (deterministic, applied in this order):

        1. **Styled headings win.** If any paragraph uses a ``Heading N``
           style, return an empty mapping — inference is fully disabled and
           styled documents parse byte-identically to previous behavior.
        2. **Candidate filter.** A paragraph may be a heading only if it passes
           all structural checks (:meth:`_heading_signal_features`): short
           (<= ``MAX_INFERRED_HEADING_LENGTH`` chars), does not end with
           sentence punctuation, contains at least one letter, and is not a
           list item (no numbering XML, no list style, no bullet prefix).
        3. **Emphasis requirement.** A structural candidate qualifies only if
           it also carries at least one emphasis signal: dominantly bold runs,
           an explicit font size larger than the body text, or ALL-CAPS text.
           A lone short paragraph is never a heading.
        4. **Level assignment precedence** (first match wins):
           a. numbering-prefix depth: ``"2.3.1 Title"`` -> level 3;
           b. font-size tier: distinct qualifying sizes sorted descending map
              to levels 1, 2, ...;
           c. bold-only / caps-only at body size -> one level deeper than the
              deepest font-size tier (level 1 when there are no tiers).
        5. **Ratio guard.** If more than ``MAX_CANDIDATE_RATIO`` of the
           document's paragraphs qualify, the signals are deemed unreliable
           (e.g. a fully bold document) and nothing is inferred.

        Returns:
            Mapping from paragraph XML element (``CT_P``) to inferred level.
        """
        paragraphs = doc.paragraphs
        if any(p.style.name.startswith('Heading') for p in paragraphs):
            return {}

        candidates: list[dict[str, Any]] = []
        body_sizes: list[float] = []
        total_non_empty = 0
        for para in paragraphs:
            text = para.text.strip()
            if not text:
                continue
            total_non_empty += 1
            features = self._heading_signal_features(para, text)
            if features is None:
                body_sizes.extend(self._explicit_run_sizes(para))
            else:
                candidates.append(features)

        if not candidates:
            return {}

        body_size = self._body_font_size(doc, body_sizes)

        qualified = [
            f for f in candidates
            if f['bold'] or f['all_caps'] or
            (f['size'] is not None and f['size'] > body_size + 0.1)
        ]
        if not qualified:
            return {}
        if (total_non_empty >= self.MIN_PARAGRAPHS_FOR_RATIO_GUARD and
                len(qualified) > total_non_empty * self.MAX_CANDIDATE_RATIO):
            logger.debug(
                "Heading inference aborted: %d of %d paragraphs look like headings",
                len(qualified), total_non_empty,
            )
            return {}

        tier_sizes = sorted(
            {f['size'] for f in qualified
             if f['size'] is not None and f['size'] > body_size + 0.1},
            reverse=True,
        )
        level_by_size = {size: index + 1 for index, size in enumerate(tier_sizes)}
        default_level = min(len(tier_sizes) + 1, self.MAX_HEADING_LEVEL)

        inferred_levels: dict[CT_P, int] = {}
        for features in qualified:
            if features['numbering_depth'] > 0:
                level = min(features['numbering_depth'], self.MAX_HEADING_LEVEL)
            elif features['size'] in level_by_size:
                level = level_by_size[features['size']]
            else:
                level = default_level
            inferred_levels[features['element']] = level
        return inferred_levels

    def _heading_signal_features(self, para: Paragraph, text: str) -> dict[str, Any] | None:
        """Return heading signal features for a paragraph, or None if the
        paragraph fails the structural checks and can never be a heading."""
        if len(text) > self.MAX_INFERRED_HEADING_LENGTH:
            return None
        if text.endswith(self.SENTENCE_END_CHARS):
            return None
        if not any(char.isalpha() for char in text):
            return None
        if text.startswith(self.BULLET_PREFIXES):
            return None
        # List items (explicit numbering XML or list styles) are never headings.
        if para._p.pPr is not None and para._p.pPr.numPr is not None:
            return None
        style_name_lower = para.style.name.lower()
        if 'list' in style_name_lower or 'bullet' in style_name_lower or 'number' in style_name_lower:
            return None

        letters = [char for char in text if char.isalpha()]
        all_caps = (
            len(letters) >= self.MIN_ALL_CAPS_LETTERS
            and all(char.isupper() for char in letters)
        )
        sizes = self._explicit_run_sizes(para)
        numbering_match = self.NUMBERING_PREFIX_RE.match(text)
        numbering_depth = numbering_match.group(1).count('.') + 1 if numbering_match else 0

        return {
            'element': para._p,
            'bold': self._is_dominantly_bold(para),
            'all_caps': all_caps,
            'size': max(sizes) if sizes else None,
            'numbering_depth': numbering_depth,
        }

    def _is_dominantly_bold(self, para: Paragraph) -> bool:
        """True when at least MIN_BOLD_RATIO of the run characters are bold."""
        total_chars = 0
        bold_chars = 0
        for run in para.runs:
            run_chars = len(run.text.strip())
            if run_chars == 0:
                continue
            total_chars += run_chars
            if run.bold:
                bold_chars += run_chars
        return total_chars > 0 and (bold_chars / total_chars) >= self.MIN_BOLD_RATIO

    def _explicit_run_sizes(self, para: Paragraph) -> list[float]:
        """Explicit font sizes (in points) of the paragraph's non-empty runs."""
        return [
            run.font.size.pt
            for run in para.runs
            if run.font.size is not None and run.text.strip()
        ]

    def _body_font_size(self, doc: docx.document.Document, body_sizes: list[float]) -> float:
        """Most common explicit body-text font size, with sensible fallbacks."""
        if body_sizes:
            return Counter(body_sizes).most_common(1)[0][0]
        try:
            normal_size = doc.styles['Normal'].font.size
            if normal_size is not None:
                return normal_size.pt
        except KeyError:
            pass
        return self.DEFAULT_BODY_FONT_SIZE_PT
    
    def _parse_content_elements(self, document_object: docx.document.Document) -> list[dict[str, Any]]:
        """Parses a sequence of XML elements and reconstructs them into a hierarchical list."""
        root_nodes: list[dict[str, Any]] = []
        parent_stack: list[dict[str, Any]] = []
        xml_element_iterator = document_object.element.body.iterchildren()
        for element in xml_element_iterator:
            element_data: dict[str, Any] | None = None
            if isinstance(element, CT_P):
                para = self._find_paragraph(document_object, element)
                if para and para.text and para.text.strip():
                    element_data = self._process_paragraph(para)
            elif isinstance(element, CT_Tbl):
                table = self._find_table(document_object, element)
                if table:
                    element_data = self._process_table(table)
            else:
                logger.debug("Skipping unsupported element type: %s", type(element))
                continue

            if not element_data:
                continue

            # Create the node for the current element, adding a 'children' list
            node = {**element_data, 'children': []}

            if node['type'] == 'heading':
                while parent_stack and \
                        ((parent_stack[-1]['type'] == 'heading' and parent_stack[-1]['level'] >= node['level']) or \
                        (parent_stack[-1]['type'] in ['list_container', 'list_item'])):
                    parent_stack.pop()

                if not parent_stack:
                    root_nodes.append(node)
                else:
                    parent_stack[-1]['children'].append(node)
                parent_stack.append(node)

            elif node['type'] == 'list_item':
                li_level = node['level']
                li_num_id = node['num_id']

                while parent_stack:
                    p_on_stack = parent_stack[-1]
                    if p_on_stack['type'] == 'list_container':
                        if p_on_stack['num_id'] == li_num_id:
                            if p_on_stack['level'] == li_level: break
                            elif p_on_stack['level'] < li_level: break
                            else: parent_stack.pop()
                        else: parent_stack.pop()
                    elif p_on_stack['type'] == 'list_item':
                        if p_on_stack['num_id'] == li_num_id and li_level > p_on_stack['level']: break
                        else: parent_stack.pop()
                    elif p_on_stack['type'] == 'heading': break
                    else:
                        parent_stack.pop()

                current_parent_on_stack = parent_stack[-1] if parent_stack else None

                if current_parent_on_stack and \
                    current_parent_on_stack['type'] == 'list_container' and \
                    current_parent_on_stack['num_id'] == li_num_id and \
                    current_parent_on_stack['level'] == li_level:
                    current_parent_on_stack['children'].append(node)
                    parent_stack.append(node)
                else:
                    list_container_node = {
                        'type': 'list_container',
                        'level': li_level,
                        'num_id': li_num_id,
                        'children': [node]
                    }
                    if not current_parent_on_stack:
                        root_nodes.append(list_container_node)
                    else:
                        current_parent_on_stack['children'].append(list_container_node)
                    parent_stack.append(list_container_node)
                    parent_stack.append(node)

            elif node['type'] in ['paragraph', 'table']:
                while parent_stack and parent_stack[-1]['type'] in ['list_item', 'list_container']:
                    parent_stack.pop()
                
                if not parent_stack:
                    root_nodes.append(node)
                else:
                    parent_stack[-1]['children'].append(node)

        return root_nodes

    def _find_paragraph(self, doc: docx.document.Document, element: CT_P) -> Paragraph:
        """Find paragraph object by XML element"""
        for para in doc.paragraphs:
            if para._element == element:
                return para
        raise ValueError("Paragraph not found")

    def _find_table(self, doc: docx.document.Document, element: CT_Tbl) -> Table:
        """Find table object by XML element"""
        for table in doc.tables:
            if table._element == element:
                return table
        raise ValueError("Table not found")

    def _process_paragraph(self, para: Paragraph) -> dict[str, Any]:
        """Process a paragraph into an element dictionary with type, content, level, and num_id for lists."""
        text = para.text.strip()

        # Heading (styled) - always takes precedence
        if para.style.name.startswith('Heading'):
            level_str = para.style.name.replace('Heading', '').strip() or '1'
            heading_level = int(level_str)
            self.current_heading_level = heading_level  # Update current heading level
            return {
                "type": "heading",
                "level": heading_level,
                "content": text
            }

        # Heading (inferred from formatting; only populated for documents
        # without any styled headings - see _infer_heading_levels)
        inferred_level = self._inferred_heading_levels.get(para._p)
        if inferred_level is not None:
            self.current_heading_level = inferred_level
            return {
                "type": "heading",
                "level": inferred_level,
                "content": text
            }

        # List item from oxml (most reliable)
        if para._p.pPr is not None and para._p.pPr.numPr is not None:
            num_pr = para._p.pPr.numPr
            # ilvl (indentation level) is 0-indexed
            ilvl = num_pr.ilvl.val if num_pr.ilvl is not None and num_pr.ilvl.val is not None else 0
            # numId references the numbering definition
            num_id = num_pr.numId.val if num_pr.numId is not None and num_pr.numId.val is not None else 0 
            
            return {
                "type": "list_item",
                "level": ilvl,  # Use ilvl as the 'level' for list items
                "content": text,
                "num_id": num_id
            }

        # Fallback: Style-based list detection
        style_name_lower = para.style.name.lower()
        if 'list' in style_name_lower or 'bullet' in style_name_lower or 'number' in style_name_lower:
            return {
                "type": "list_item",
                "level": 0,  # Default ilvl if unknown
                "content": text,
                "num_id": -1 # Indicate unknown num_id for fallback
            }

        # Fallback: Text-based list detection
        if text.startswith(('- ', '• ', '* ')) or \
           (text.split('.', 1)[0].isdigit() and len(text.split('.', 1)[0]) < 3 and '.' in text):
            return {
                "type": "list_item",
                "level": 0,  # Default ilvl if unknown
                "content": text,
                "num_id": -1 # Indicate unknown num_id for fallback
            }

        # Paragraph
        return {
            "type": "paragraph",
            "level": self.current_heading_level if self.current_heading_level > 0 else 0,
            "content": text
        }

    def _process_table(self, table) -> dict[str, Any]:
        """
        Process a table into an element dictionary.
        The first row is assumed to be the header.
        Subsequent rows are stored individually.
        """
        header_cells: list[str] = []
        data_rows_content: list[list[str]] = []

        if table.rows:
            first_row_cells = table.rows[0].cells
            for cell in first_row_cells:
                cell_para_texts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
                header_cells.append(" ".join(cell_para_texts))

            # Process subsequent rows as data rows
            for i in range(1, len(table.rows)):
                row = table.rows[i]
                current_row_cells_text: list[str] = []
                for cell in row.cells:
                    cell_para_texts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
                    current_row_cells_text.append(" ".join(cell_para_texts))
                if any(current_row_cells_text): # Add row only if it has some content
                    data_rows_content.append(current_row_cells_text)

        # The 'content' field is removed in favor of structured header/rows.
        # If a single string representation is still needed elsewhere, 
        # it would need to be constructed by the consumer of this structure.
        return {
            "type": "table",
            "level": self.current_heading_level if self.current_heading_level > 0 else 0,
            "num_rows": len(table.rows), # Total rows including potential header
            "num_cols": len(table.columns) if table.columns else 0,
            "header": header_cells,
            "data_rows": data_rows_content # List of lists, where each inner list contains cell strings for a data row
        }
