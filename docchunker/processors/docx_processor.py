"""
Processor for DOCX documents.
"""

import os
import re
import uuid
from typing import Dict, List, Optional, Tuple, Union

import docx
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from docchunker.models.chunk import Chunk
from docchunker.models.document import Document
from docchunker.processors.base import BaseProcessor


class DocxProcessor(BaseProcessor):
    """
    Processor for DOCX documents.
    
    This class handles the extraction and chunking of content from DOCX files,
    with special handling for complex structures like tables and lists.
    """
    
    def process(self, file_path: str) -> Document:
        """
        Process a DOCX document and extract its content and structure.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            A Document object containing the extracted content
            
        Raises:
            ValueError: If the file cannot be processed
            FileNotFoundError: If the file doesn't exist
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not file_path.lower().endswith(".docx"):
            raise ValueError(f"File is not a DOCX document: {file_path}")
        
        document_id = self._generate_document_id(file_path)
        
        try:
            docx_document = docx.Document(file_path)
            metadata = self._extract_metadata(docx_document)
            
            # Create our document model
            document = Document(
                document_id=document_id,
                file_path=file_path,
                metadata={
                    "title": os.path.basename(file_path),
                    "file_format": "docx",
                    **metadata
                }
            )
            
            # Process the document content
            self._process_content(docx_document, document)
            
            return document
        
        except Exception as e:
            raise ValueError(f"Failed to process DOCX document: {str(e)}")
    
    def chunk(self, document: Document) -> List[Chunk]:
        """
        Chunk a processed document into smaller, semantically meaningful parts.
        
        Args:
            document: The processed document to chunk
            
        Returns:
            A list of Chunk objects
        """
        # The document already contains initial chunks from the processing stage
        # Here we apply additional chunking logic based on semantics and size constraints
        
        result_chunks = []
        current_section_chunks = []
        current_section_text = ""
        current_section_type = None
        
        for chunk in document.chunks:
            chunk_type = chunk.metadata.get("type")
            
            # Handle special chunk types that should be preserved as-is
            if chunk_type in [Chunk.TYPE_TABLE, Chunk.TYPE_NESTED_TABLE, Chunk.TYPE_IMAGE, 
                             Chunk.TYPE_MERGED_CELL, Chunk.TYPE_EQUATION, Chunk.TYPE_CODE]:
                # First, add any accumulated section chunks
                if current_section_text:
                    section_chunks = self._chunk_text(
                        current_section_text, 
                        document.document_id, 
                        current_section_type or Chunk.TYPE_TEXT
                    )
                    result_chunks.extend(section_chunks)
                    current_section_text = ""
                    current_section_chunks = []
                    current_section_type = None
                
                # Then add the special chunk as-is
                result_chunks.append(chunk)
                continue
            
            # For heading chunks, they represent section boundaries
            if chunk_type == Chunk.TYPE_HEADING:
                # End previous section
                if current_section_text:
                    section_chunks = self._chunk_text(
                        current_section_text, 
                        document.document_id, 
                        current_section_type or Chunk.TYPE_TEXT
                    )
                    result_chunks.extend(section_chunks)
                
                # Start new section with the heading
                result_chunks.append(chunk)
                current_section_text = ""
                current_section_chunks = []
                current_section_type = Chunk.TYPE_TEXT
                continue
            
            # For regular text, lists, etc., accumulate into the current section
            if not current_section_type:
                current_section_type = chunk_type
            
            # If chunk would make section too large, chunk what we have so far
            if current_section_text and len(current_section_text) + len(chunk.text) > self.chunk_size * 2:
                section_chunks = self._chunk_text(
                    current_section_text, 
                    document.document_id, 
                    current_section_type
                )
                result_chunks.extend(section_chunks)
                current_section_text = chunk.text
                current_section_chunks = [chunk]
            else:
                if current_section_text:
                    current_section_text += "\n\n"
                current_section_text += chunk.text
                current_section_chunks.append(chunk)
        
        # Add any remaining section chunks
        if current_section_text:
            section_chunks = self._chunk_text(
                current_section_text, 
                document.document_id, 
                current_section_type or Chunk.TYPE_TEXT
            )
            result_chunks.extend(section_chunks)
        
        return result_chunks
    
    def _extract_metadata(self, docx_document: DocxDocument) -> Dict[str, str]:
        """
        Extract metadata from a DOCX document.
        
        Args:
            docx_document: The python-docx Document object
            
        Returns:
            Dictionary of metadata
        """
        metadata = {}
        
        # Extract core properties
        core_properties = docx_document.core_properties
        
        if core_properties.title:
            metadata["title"] = core_properties.title
        
        if core_properties.author:
            metadata["author"] = core_properties.author
        
        if core_properties.created:
            metadata["created_date"] = str(core_properties.created)
        
        if core_properties.modified:
            metadata["modified_date"] = str(core_properties.modified)
        
        # Count pages (approximate since python-docx doesn't provide direct page count)
        # This is a rough estimate based on paragraph count
        paragraph_count = len(docx_document.paragraphs)
        estimated_page_count = max(1, paragraph_count // 40)  # Assuming ~40 paragraphs per page
        metadata["page_count"] = estimated_page_count
        
        return metadata
    
    def _process_content(self, docx_document: DocxDocument, document: Document) -> None:
        """
        Process the content of a DOCX document.
        
        Args:
            docx_document: The python-docx Document object
            document: Our document model to populate
        """
        # Process all elements in the document body
        for element in docx_document.element.body:
            if isinstance(element, CT_P):
                # It's a paragraph
                paragraph = Paragraph(element, docx_document)
                self._process_paragraph(paragraph, document)
            
            elif isinstance(element, CT_Tbl):
                # It's a table
                table = Table(element, docx_document)
                self._process_table(table, document)
        
        # Process headers and footers if needed
        # Note: This is more complex and would need additional code
    
    def _process_paragraph(self, paragraph: Paragraph, document: Document) -> None:
        """
        Process a paragraph element.
        
        Args:
            paragraph: The paragraph to process
            document: Our document model to populate
        """
        text = paragraph.text.strip()
        
        if not text:
            return
        
        # Check if this is a heading
        if paragraph.style.name.startswith('Heading'):
            try:
                level = int(paragraph.style.name.replace('Heading', '').strip())
            except ValueError:
                level = 1  # Default to level 1 if we can't parse the level
            
            chunk = Chunk.create_heading_chunk(
                text=text,
                document_id=document.document_id,
                level=level
            )
            document.chunks.append(chunk)
            return
        
        # Check if this is a list item
        if self.handle_lists and self._is_list_item(paragraph):
            list_info = self._extract_list_info(paragraph)
            chunk = Chunk.create_list_chunk(
                text=text,
                document_id=document.document_id,
                list_type=list_info["list_type"],
                depth=list_info["depth"]
            )
            document.chunks.append(chunk)
            return
        
        # Regular paragraph
        chunk = Chunk.create_text_chunk(
            text=text,
            document_id=document.document_id
        )
        document.chunks.append(chunk)
    
    def _process_table(self, table: Table, document: Document, parent_table_id: Optional[str] = None) -> None:
        """
        Process a table element.
        
        Args:
            table: The table to process
            document: Our document model to populate
            parent_table_id: ID of the parent table if this is a nested table
        """
        if not self.handle_tables:
            # If not handling tables specially, just extract text
            table_text = self._extract_table_text(table)
            chunk = Chunk.create_text_chunk(
                text=table_text,
                document_id=document.document_id
            )
            document.chunks.append(chunk)
            return
        
        # Generate a unique ID for this table
        table_id = str(uuid.uuid4())[:8]
        
        # Check if this table has merged cells
        has_merged_cells = self._table_has_merged_cells(table)
        
        # Check for header row (simplistic check)
        has_header = False
        if len(table.rows) > 1:
            # If first row has different formatting, consider it a header
            if table.rows[0].cells[0]._element.tcPr is not None:
                has_header = True
        
        # Create the table chunk
        if parent_table_id:
            # This is a nested table
            chunk = Chunk.create_nested_table_chunk(
                text=self._extract_table_text(table),
                document_id=document.document_id,
                table_id=table_id,
                parent_table_id=parent_table_id,
                rows=len(table.rows),
                columns=len(table.columns)
            )
        else:
            # This is a top-level table
            chunk = Chunk.create_table_chunk(
                text=self._extract_table_text(table),
                document_id=document.document_id,
                table_id=table_id,
                rows=len(table.rows),
                columns=len(table.columns),
                has_header=has_header,
                has_merged_cells=has_merged_cells
            )
        
        document.chunks.append(chunk)
        
        # Process individual cells
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                # Skip cells that are part of a merge and not the main cell
                if self._is_merged_continuation_cell(table, i, j):
                    continue
                
                # Check if this is a merged cell
                is_merged = False
                row_span = 1
                col_span = 1
                
                if has_merged_cells:
                    spans = self._get_cell_spans(table, i, j)
                    row_span = spans[0]
                    col_span = spans[1]
                    is_merged = row_span > 1 or col_span > 1
                
                # Process cell content
                cell_paragraphs = cell.paragraphs
                cell_text = "\n".join([p.text for p in cell_paragraphs if p.text.strip()])
                
                if not cell_text.strip():
                    continue
                
                # Check for nested tables in this cell
                nested_tables = cell._element.xpath('.//w:tbl')
                if nested_tables:
                    for nested_table_elem in nested_tables:
                        nested_table = Table(nested_table_elem, cell._parent._parent)
                        self._process_table(nested_table, document, table_id)
                
                # Create cell chunk
                if is_merged:
                    chunk = Chunk.create_merged_cell_chunk(
                        text=cell_text,
                        document_id=document.document_id,
                        table_id=table_id,
                        row_span=row_span,
                        col_span=col_span,
                        start_row=i,
                        start_col=j
                    )
                else:
                    chunk = Chunk(
                        text=cell_text,
                        metadata={
                            "type": Chunk.TYPE_TABLE_CELL,
                            "document_id": document.document_id,
                            "table_id": table_id,
                            "row": i,
                            "column": j,
                            "is_header": (i == 0 and has_header)
                        }
                    )
                
                document.chunks.append(chunk)
    
    def _extract_table_text(self, table: Table) -> str:
        """
        Extract text from a table in a formatted way.
        
        Args:
            table: The table to extract text from
            
        Returns:
            Formatted table text
        """
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                text = " ".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                cells.append(text)
            rows.append(" | ".join(cells))
        
        return "\n".join(rows)
    
    def _is_list_item(self, paragraph: Paragraph) -> bool:
        """
        Check if a paragraph is a list item.
        
        Args:
            paragraph: The paragraph to check
            
        Returns:
            True if the paragraph is a list item, False otherwise
        """
        # Check for numbering
        if paragraph._element.pPr and paragraph._element.pPr.numPr:
            return True
        
        # Check for bullet-like markers
        text = paragraph.text.strip()
        bullet_patterns = [
            r'^\s*[\•\-\*\>\◦\▪\■\○\●]\s',
            r'^\s*[a-zA-Z0-9]\.\s',
            r'^\s*[(][a-zA-Z0-9][)]\s',
        ]
        
        for pattern in bullet_patterns:
            if re.match(pattern, text):
                return True
        
        return False
    
    def _extract_list_info(self, paragraph: Paragraph) -> Dict[str, Union[str, int]]:
        """
        Extract information about a list item.
        
        Args:
            paragraph: The list item paragraph
            
        Returns:
            Dictionary with list type and depth
        """
        # Default values
        info = {
            "list_type": "bullet",
            "depth": 0
        }
        
        # Check for numbering
        if paragraph._element.pPr and paragraph._element.pPr.numPr:
            # Try to determine if numbered or bulleted
            if paragraph._element.pPr.numPr.numId:
                # This is a simplified approach - a more thorough implementation
                # would check the actual numbering definition
                info["list_type"] = "numbered"
            
            # Try to determine depth
            if paragraph._element.pPr.numPr.ilvl:
                ilvl = paragraph._element.pPr.numPr.ilvl.val
                if ilvl is not None:
                    info["depth"] = int(ilvl)
            
            return info
        
        # Fallback: Check text patterns
        text = paragraph.text.strip()
        indent = len(text) - len(text.lstrip())
        info["depth"] = min(3, indent // 2)  # Approximation based on indent
        
        # Try to determine type from text pattern
        if re.match(r'^\s*[0-9]+\.', text) or re.match(r'^\s*[ivxlcdm]+\.', text) or re.match(r'^\s*[a-zA-Z]\.', text):
            info["list_type"] = "numbered"
        
        return info
    
    def _table_has_merged_cells(self, table: Table) -> bool:
        """
        Check if a table has merged cells.
        
        Args:
            table: The table to check
            
        Returns:
            True if the table has merged cells, False otherwise
        """
        for row in table.rows:
            for cell in row.cells:
                # Check for vertical merge (rowspan)
                if cell._element.tcPr and cell._element.tcPr.vMerge:
                    return True
                
                # Check for horizontal merge (colspan)
                if cell._element.tcPr and cell._element.tcPr.gridSpan and cell._element.tcPr.gridSpan.val:
                    if int(cell._element.tcPr.gridSpan.val) > 1:
                        return True
        
        return False
    
    def _is_merged_continuation_cell(self, table: Table, row_idx: int, col_idx: int) -> bool:
        """
        Check if a cell is a continuation cell in a merge (not the main cell).
        
        Args:
            table: The table containing the cell
            row_idx: Row index of the cell
            col_idx: Column index of the cell
            
        Returns:
            True if this is a continuation cell, False otherwise
        """
        try:
            cell = table.rows[row_idx].cells[col_idx]
            
            # Check for vertical merge continuation
            if cell._element.tcPr and cell._element.tcPr.vMerge:
                vMerge = cell._element.tcPr.vMerge
                if vMerge.val is None or vMerge.val == "continue":
                    return True
            
            # For horizontal merges, we rely on the fact that merged cells 
            # are only represented by their starting cell in the python-docx model
            
            return False
        except IndexError:
            return False
    
    def _get_cell_spans(self, table: Table, row_idx: int, col_idx: int) -> Tuple[int, int]:
        """
        Get the row and column spans for a cell.
        
        Args:
            table: The table containing the cell
            row_idx: Row index of the cell
            col_idx: Column index of the cell
            
        Returns:
            Tuple of (row_span, col_span)
        """
        try:
            cell = table.rows[row_idx].cells[col_idx]
            row_span = 1
            col_span = 1
            
            # Get column span (gridSpan)
            if cell._element.tcPr and cell._element.tcPr.gridSpan and cell._element.tcPr.gridSpan.val:
                col_span = int(cell._element.tcPr.gridSpan.val)
            
            # Get row span (vMerge)
            if cell._element.tcPr and cell._element.tcPr.vMerge:
                vMerge = cell._element.tcPr.vMerge
                if vMerge.val == "restart" or vMerge.val == 1:
                    # This is the start of a vertical merge
                    # Count how many cells below are merged with this one
                    for i in range(row_idx + 1, len(table.rows)):
                        try:
                            below_cell = table.rows[i].cells[col_idx]
                            if below_cell._element.tcPr and below_cell._element.tcPr.vMerge:
                                below_vMerge = below_cell._element.tcPr.vMerge
                                if below_vMerge.val is None or below_vMerge.val == "continue":
                                    row_span += 1
                                else:
                                    break
                            else:
                                break
                        except IndexError:
                            break
            
            return (row_span, col_span)
        
        except IndexError:
            return (1, 1)